"""DOCX rendering for Scribe reports.

Kept separate from render.py because it needs the optional `python-docx`
dependency. `render_docx` raises ImportError if it isn't installed — callers
should catch and surface a clear "DOCX export unavailable" message.

Builds the Word document from the same report `data` dict the markdown/HTML
renderers use (NOT by converting HTML), so the structure is native and clean.
"""

from __future__ import annotations

import io

from .labels import Labels
from .render import (
    TEMPLATE_DEFAULTS,
    _ev_fields,
    _manifest_ai_line,
    _ts,
)


def render_docx(data: dict, tpl: dict | None = None, language: str | None = None) -> bytes:
    """Render the report as a .docx and return the raw bytes.

    Raises ImportError when python-docx is missing.
    """
    from docx import Document  # type: ignore
    from docx.shared import Pt

    tpl = tpl or dict(TEMPLATE_DEFAULTS)
    L = Labels(language)
    sections = tpl.get("sections") or {}
    case = data.get("case") or {}
    name = case.get("name", case.get("case_id", "Case"))

    title_prefix = tpl.get("title_prefix")
    if not title_prefix or title_prefix == TEMPLATE_DEFAULTS["title_prefix"]:
        title_prefix = L("title_prefix")

    doc = Document()
    doc.add_heading(f"{title_prefix} — {name}", level=0)

    meta_bits = [f"{L('generated')} {_now_label()}"]
    if case.get("company"):
        meta_bits.append(f"{L('company')}: {case['company']}")
    doc.add_paragraph(" · ".join(meta_bits)).italic = True

    def h2(text):
        doc.add_heading(text, level=1)

    def h3(text):
        doc.add_heading(text, level=2)

    def bullet(text):
        doc.add_paragraph(str(text), style="List Bullet")

    # ── Manifest
    manifest = data.get("manifest") or {}
    if sections.get("manifest", True) and manifest:
        h2(L("manifest_title"))
        doc.add_paragraph(L("manifest_blurb"))
        if manifest.get("total_events"):
            bullet(L("m_events", n=f"{int(manifest['total_events']):,}"))
        bullet(L("m_flagged", n=manifest.get("flagged_count", 0)))
        if manifest.get("pinned_count"):
            bullet(L("m_pinned", n=manifest["pinned_count"]))
        bullet(L("m_modules", hit=manifest.get("module_hit_run_count", 0), total=manifest.get("module_run_count", 0)))
        if manifest.get("saved_search_count"):
            bullet(L("m_saved", n=manifest["saved_search_count"]))
        if manifest.get("killchain_count"):
            bullet(L("m_killchains", n=manifest["killchain_count"]))
        bullet(_manifest_ai_line(L, manifest))

    # ── Executive summary
    agg = data.get("aggregates") or {}
    flagged = data.get("flagged") or []
    pinned = data.get("pinned") or []
    wl = data.get("watchlist") or {}
    # True totals (ES track_total_hits) — `pinned`/`flagged` are bounded samples,
    # so len() would cap the reported counts at the sample size.
    flagged_total = data.get("flagged_count", len(flagged))
    if sections.get("exec_summary", True):
        h2(L("exec_summary"))
        if agg.get("total_events"):
            bullet(L("total_events", n=f"{int(agg['total_events']):,}"))
        bullet(L("flagged_review", n=flagged_total))
        if agg.get("cti"):
            bullet(L("cti_matched", n=len(agg["cti"])))
        if wl.get("hits"):
            bullet(L("wl_hits", n=len(wl["hits"])))

    # ── Activity overview (aggregates as plain "value — count" bullets)
    if sections.get("overview", True) and agg:
        overview_rows = [
            (L("events_over_time"), agg.get("timeline") or [], L("u_events")),
            (L("artifact_types"), agg.get("artifact_types") or [], L("u_events")),
            (L("top_src_ips"), agg.get("top_src_ips") or [], L("u_events")),
        ]
        if any(items for _, items, _ in overview_rows):
            h2(L("activity_overview"))
            for title, items, unit in overview_rows:
                if not items:
                    continue
                h3(title)
                for it in items[:15]:
                    bullet(f"{it.get('value')} — {int(it.get('count', 0)):,} {unit}")

    # ── Module analysis
    modules = data.get("modules") or []
    if sections.get("modules", True) and modules:
        h2(L("module_analysis"))
        for r in modules[:40]:
            lv = r.get("hits_by_level") or {}
            sev = ", ".join(f"{int(n)} {k}" for k, n in lv.items() if n)
            line = f"{r.get('module_id', '?')} ({r.get('status', '')}) — {r.get('total_hits', 0)} {L('u_hits')}"
            if sev:
                line += f" [{sev}]"
            bullet(line)

    # ── Findings (unified findings store)
    findings = data.get("findings") or {}
    f_items = findings.get("items") or []
    if sections.get("findings", True) and f_items:
        h2("Findings")
        doc.add_paragraph(
            "The unified findings store — every analysis surface's saved output "
            "(IOCs, anomalies, MITRE coverage, kill chains, modules, co-pilot, "
            f"manual). {findings.get('total', len(f_items))} total."
        ).italic = True
        by_kind = findings.get("by_kind") or {}
        by_sev = findings.get("by_severity") or {}
        if by_kind:
            bullet("by kind: " + ", ".join(f"{k} ({n})" for k, n in by_kind.items()))
        if by_sev:
            bullet("by severity: " + ", ".join(f"{k} ({n})" for k, n in by_sev.items()))
        grouped: dict[str, list] = {}
        for it in f_items:
            grouped.setdefault(it.get("kind", "other"), []).append(it)
        for kind, rows in grouped.items():
            h3(f"{kind} ({len(rows)})")
            for it in rows[:50]:
                sev = it.get("severity", "informational")
                ts = _ts(it.get("timestamp", ""))
                title = (it.get("message") or (it.get("finding") or {}).get("title") or "").strip()
                src = it.get("source_feature", "")
                line = f"[{sev}] " + (f"{ts} " if ts else "") + title
                if src:
                    line += f" (via {src})"
                bullet(line)
            if len(rows) > 50:
                doc.add_paragraph(L("and_more", n=len(rows) - 50)).italic = True

    # ── Saved searches
    saved = data.get("saved_searches") or []
    if sections.get("saved_searches", True) and saved:
        h2(L("saved_title"))
        doc.add_paragraph(L("saved_blurb"))
        for s in saved:
            bullet(f"{s.get('name', '?')} — {L('saved_matches', n=int(s.get('count', 0)))}: {s.get('query', '')}")

    # ── Correlated kill chains
    killchains = data.get("killchains") or []
    if sections.get("killchains", True) and killchains:
        h2(L("killchains_title"))
        doc.add_paragraph(L("killchains_blurb"))
        for kc in killchains:
            a = kc.get("anchor") or {}
            h3(f"{L('kc_anchor')}: {a.get('summary') or a.get('fo_id', '?')}")
            if kc.get("tactics_covered"):
                doc.add_paragraph(f"{L('kc_tactics')}: {', '.join(kc['tactics_covered'])}")
            for st in kc.get("steps", [])[:30]:
                tac = st.get("tactic") or st.get("phase") or ""
                bullet(f"{_ts(st.get('ts', ''))} — {tac} {st.get('technique', '')}: {(st.get('summary') or '')[:200]}")

    # ── Threat intel matches
    cti = agg.get("cti") or []
    if sections.get("threat_intel", True) and cti:
        h2(L("threat_intel_matches"))
        for it in cti[:60]:
            ctx = []
            if it.get("count"):
                ctx.append(f"{it['count']} {L('u_hits')}")
            if it.get("last_seen"):
                ctx.append(f"{L('last')} {_ts(it['last_seen'])}")
            if it.get("feed"):
                ctx.append(f"{L('feed')} {it['feed']}")
            if it.get("threat"):
                ctx.append(str(it["threat"]))
            if it.get("prior_cases"):
                ctx.append(L("prior_cases", n=it["prior_cases"]))
            bullet(f"{it.get('value')} ({it.get('type', 'ioc')}) — " + ", ".join(ctx))

    # ── AI narrative (already prose) — drop in as paragraphs
    ai_report = data.get("ai_report")
    if sections.get("ai_report", True) and ai_report and (ai_report.get("content") or "").strip():
        h2(L("ai_investigation"))
        for ln in ai_report["content"].splitlines():
            ln = ln.rstrip()
            if not ln:
                continue
            if ln.startswith("### "):
                h3(ln[4:])
            elif ln.startswith("## "):
                h3(ln[3:])
            elif ln.startswith("# "):
                h3(ln[2:])
            elif ln.lstrip().startswith(("- ", "* ")):
                bullet(ln.lstrip()[2:])
            else:
                doc.add_paragraph(ln)

    # ── Pinned / flagged events
    if sections.get("pinned", True) and pinned:
        h2(L("key_evidence"))
        for ev in pinned:
            f = _ev_fields(ev)
            line = f"{f['ts']} [{f['type']}] {f['host']} {f['user']} — {f['msg']}"
            if f["src"]:
                line += f" (from {f['src']})"
            bullet(line)

    max_flagged = int(tpl.get("max_flagged") or 50)
    if sections.get("flagged", True) and flagged:
        h2(L("flagged_events"))
        for ev in flagged[:max_flagged]:
            f = _ev_fields(ev)
            line = f"{f['ts']} [{f['type']}] {f['host']} {f['user']} — {f['msg']}"
            if f["src"]:
                line += f" (from {f['src']})"
            bullet(line)
        shown = min(len(flagged), max_flagged)
        if flagged_total > shown:
            doc.add_paragraph(L("and_more", n=flagged_total - shown)).italic = True

    # ── MITRE coverage
    mitre = data.get("mitre") or {}
    techs = mitre.get("techniques", []) if sections.get("mitre", True) else []
    if techs:
        h2(L("mitre_coverage"))
        by_tactic: dict[str, list] = {}
        for t in techs:
            by_tactic.setdefault(t.get("tactic") or L("unknown"), []).append(t)
        for tactic, items in sorted(by_tactic.items()):
            h3(tactic)
            for t in sorted(items, key=lambda x: -x.get("count", 0)):
                bullet(f"{t['id']} {t.get('name', '')} — {t.get('count', 0)} {L('u_event')}")

    # ── Detections fired
    det = data.get("detections") or {}
    if sections.get("detections", True) and det.get("matches"):
        h2(L("detections_fired"))
        doc.add_paragraph(f"{L('last_run')}: {_ts(det.get('ran_at', ''))}").italic = True
        for m in det["matches"]:
            rule = m.get("rule") or {}
            bullet(f"{rule.get('name', '?')} ({m.get('match_count', 0)} {L('u_matches')})")
            if rule.get("description"):
                doc.add_paragraph(rule["description"])

    # ── Watchlist hits
    if sections.get("watchlist", True) and wl.get("hits"):
        h2(L("watchlist_hits"))
        doc.add_paragraph(f"{L('auto_sweep')}: {_ts(wl.get('ran_at', ''))}").italic = True
        for h in wl["hits"]:
            bullet(
                f"{h.get('label')} ({h.get('kind')}) — "
                f"{h.get('hits')} {L('u_matches')}: {h.get('query')}"
            )

    # ── Analyst notes
    notes = data.get("notes") or ""
    if sections.get("notes", True) and notes:
        h2(L("analyst_notes"))
        for ln in notes.splitlines():
            if ln.strip():
                doc.add_paragraph(ln)

    if (tpl.get("footer_md") or "").strip():
        p = doc.add_paragraph(tpl["footer_md"].strip())
        for run in p.runs:
            run.font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _now_label() -> str:
    from datetime import UTC, datetime

    return _ts(datetime.now(UTC).isoformat())
